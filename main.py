import pandas as pd
import numpy as np
import nltk
from docx import Document

#df = pd.read_csv("HH_CV_database_40K_str.csv")
document = Document('Resume.docx')
core_properties = document.core_properties
print(core_properties.author)
for paragraph in document.paragraphs:
  #print(paragraph.text)
  #print('99')
  pass
  
def reading_docx(path_file):
    """
    Read a docx
    Return a DataFrame with word by line
    """
    
    print ("Loading file : " + path_file.split('/')[-1])
    if path_file.split('.')[-1] == "docx":
        doc = Document(path_file)
        paragraphs = doc.paragraphs
        my_document = []
        for paragraph in paragraphs:
            my_document.append({'doc_name' : path_file.split('/')[-1],
                                'paragraph' : paragraph.text})

        documents_df = pd.DataFrame(my_document)
        documents_df.reset_index(inplace=True)
        
        word_list = []
        context_word =[]
        for idx , row in documents_df.T.iteritems():
            tokenized = nltk.word_tokenize(row['paragraph'])
            word_list.extend(tokenized)
            # Some context paragraph / doc
            for len_token in tokenized:
                context_word.append({'doc_name' : row['doc_name'],
                                     'paragraph_nb' : row['index']})
                                     
        # My bag of word
        word_df = pd.DataFrame(word_list, columns=['word'])
        
        # My context
        context_df = pd.DataFrame(context_word)
        
        # Merging context & bad of word
        word_df = pd.concat([word_df, context_df], axis=1)


        return word_df
    else:
        print("Error - Not a docx file : " + path_file.split('/')[-1] )
print(reading_docx('Resume.docx'))